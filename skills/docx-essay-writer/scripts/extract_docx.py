#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取docx文件内容，按格式差异换行，并添加位置标记

用法:
    python extract_docx.py <docx文件路径> [输出文件路径]

示例:
    python extract_docx.py input.docx output.md
    python extract_docx.py input.docx  # 默认输出到同目录的.md文件

注意: 此脚本会在markdown中保存详细的位置和格式信息，用于后续还原
"""
import sys
import os
from docx import Document


def get_paragraph_format_info(para):
    """获取段落的格式信息"""
    format_info = {}

    if para.paragraph_format:
        if para.paragraph_format.page_break_before:
            format_info['page_break_before'] = True
        if para.paragraph_format.keep_together:
            format_info['keep_together'] = True
        if para.paragraph_format.keep_with_next:
            format_info['keep_with_next'] = True

    return format_info


def get_run_format_info(run):
    """获取run的格式信息"""
    if run is None:
        return {}

    format_info = {
        'font_name': run.font.name if run.font.name else None,
        'font_size': run.font.size.pt if run.font.size else None,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        'strike': run.font.strike,
        'superscript': run.font.superscript,
        'subscript': run.font.subscript,
    }

    return {k: v for k, v in format_info.items() if v is not None}


def formats_equal(format1, format2):
    """比较两个格式是否相同"""
    if not format1 and not format2:
        return True
    if not format1 or not format2:
        return False

    # 比较所有格式属性，但忽略font_name为None的情况
    for key in format1.keys():
        val1 = format1.get(key)
        val2 = format2.get(key)

        if key == 'font_name':
            if val1 is not None and val2 is not None and val1 != val2:
                return False
        else:
            if val1 != val2:
                return False

    return True


def format_to_string(format_info):
    """将格式信息转换为字符串"""
    parts = []
    for key, value in format_info.items():
        if value is not None and value != '':
            parts.append(f'{key}={value}')
    return '|'.join(parts)


def get_field_code_text(run):
    """获取run中的域代码文本（instrText）"""
    field_text = []
    for child in run._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'instrText' and child.text:
            field_text.append(child.text)
    return ''.join(field_text) if field_text else None


def extract_runs_with_format(runs):
    """提取runs中的文本，按格式差异分割"""
    result = []
    current_format = None
    current_text = []

    for run in runs:
        if run.text == '':
            continue

        run_format = get_run_format_info(run)

        if current_format is None:
            current_format = run_format
            current_text.append(run.text)
        elif formats_equal(current_format, run_format):
            current_text.append(run.text)
        else:
            if current_text:
                result.append({
                    'text': ''.join(current_text),
                    'format': current_format
                })
            current_format = run_format
            current_text = [run.text]

    if current_text:
        result.append({
            'text': ''.join(current_text),
            'format': current_format
        })

    return result


def extract_table_content(table, table_idx, output_lines, path_prefix=""):
    """递归提取表格内容（包括嵌套表格）"""
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            # 提取单元格段落
            for para_idx_in_cell, para in enumerate(cell.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    # 提取文本或域代码内容
                    text_content = run.text
                    field_code = get_field_code_text(run)

                    # 如果run.text为空但有域代码，则提取域代码
                    if not text_content and field_code:
                        text_content = f'[FIELD:{field_code}]'

                    if not text_content:
                        continue

                    run_format = get_run_format_info(run)
                    format_str = format_to_string(run_format)
                    # 添加位置标记和格式标记
                    path = f"{path_prefix},ROW:{row_idx},CELL:{cell_idx},PARA:{para_idx_in_cell},RUN:{run_idx}"
                    output_lines.append(f'<!-- TABLE:{table_idx}{path},FORMAT:{format_str} -->{text_content}')

            # 递归处理嵌套表格
            if len(cell.tables) > 0:
                for nested_idx, nested_table in enumerate(cell.tables):
                    nested_path = f"{path_prefix},ROW:{row_idx},CELL:{cell_idx},NESTED:{nested_idx}"
                    extract_table_content(nested_table, table_idx, output_lines, nested_path)

        # 行结束后换行
        output_lines.append('')


def extract_docx_content(docx_path, output_path):
    """提取docx文件内容到txt文件，格式差异时换行，并添加位置和格式标记"""
    doc = Document(docx_path)

    output_lines = []

    # 提取段落内容
    for para_idx, para in enumerate(doc.paragraphs):
        # 获取段落格式信息
        para_format = get_paragraph_format_info(para)

        # 如果段落有page_break_before，添加标记
        if para_format.get('page_break_before'):
            output_lines.append(f'<!-- PARA:{para_idx},PAGE_BREAK_BEFORE -->')

        if not para.runs or all(run.text == '' for run in para.runs):
            # 检查是否有分页符（br type='page'）
            has_page_break = False
            for run in para.runs:
                for child in run._element:
                    from docx.oxml.ns import qn
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'br':
                        br_type = child.get(qn('w:type'))
                        if br_type == 'page':
                            has_page_break = True
                            break

            if has_page_break:
                output_lines.append(f'<!-- PARA:{para_idx},PAGE_BREAK -->')
            else:
                output_lines.append('')
            continue

        for run_idx, run in enumerate(para.runs):
            # 检查是否有分页符
            from docx.oxml.ns import qn
            for child in run._element:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag == 'br':
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        output_lines.append(f'<!-- PARA:{para_idx},RUN:{run_idx},PAGE_BREAK -->')
                        continue

            # 提取文本或域代码内容
            text_content = run.text
            field_code = get_field_code_text(run)

            # 如果run.text为空但有域代码，则提取域代码
            if not text_content and field_code:
                text_content = f'[FIELD:{field_code}]'

            if not text_content:
                continue

            run_format = get_run_format_info(run)
            format_str = format_to_string(run_format)
            # 添加位置标记和格式标记
            output_lines.append(f'<!-- PARA:{para_idx},RUN:{run_idx},FORMAT:{format_str} -->{text_content}')
        output_lines.append('')

    # 提取表格内容（包括嵌套表格）
    for table_idx, table in enumerate(doc.tables):
        # 表格标识
        output_lines.append('')
        output_lines.append('<!-- TABLE -->')
        output_lines.append('')

        # 递归提取表格内容
        extract_table_content(table, table_idx, output_lines, "")

        output_lines.append('')
        output_lines.append('<!-- TABLE_END -->')
        output_lines.append('')

    # 写入输出文件
    with open(output_path, 'w', encoding='utf-8') as f:
        for line in output_lines:
            f.write(line + '\n')

    print(f"内容已提取到: {output_path}")
    print(f"总行数: {len(output_lines)}")
    print(f"包含 {len(doc.paragraphs)} 个段落和 {len(doc.tables)} 个表格")


if __name__ == '__main__':
    # 解析命令行参数
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <docx文件路径> [输出文件路径]")
        print("示例:")
        print("  python extract_docx.py input.docx output.md")
        print("  python extract_docx.py input.docx")
        sys.exit(1)

    docx_path = sys.argv[1]

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # 如果未指定输出路径，使用与输入相同的文件名，但扩展名改为.md
        base_name = os.path.splitext(docx_path)[0]
        output_path = base_name + '.md'

    # 检查输入文件是否存在
    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在: {docx_path}")
        sys.exit(1)

    # 提取内容
    extract_docx_content(docx_path, output_path)
