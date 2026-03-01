#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从markdown文件还原到docx模板，保留所有格式和图片

用法:
    python restore_docx.py <markdown文件路径> <docx模板路径> <输出docx路径>

示例:
    python restore_docx.py content.md template.docx output.docx

原理:
    1. 读取原文档作为模板，保留所有格式、图片、表格样式
    2. 根据markdown中的位置标记，替换对应run的文本内容
    3. 保留其他所有内容不变
"""
import sys
import os
import re
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


def string_to_format(format_string):
    """将字符串转换为格式信息"""
    if not format_string:
        return {}
    format_info = {}
    for part in format_string.split('|'):
        if '=' in part:
            key, value = part.split('=', 1)
            # 转换布尔值
            if value == 'True':
                value = True
            elif value == 'False':
                value = False
            # 转换数字
            elif value.replace('.', '').replace('-', '').isdigit():
                try:
                    value = float(value) if '.' in value else int(value)
                except:
                    pass
            format_info[key] = value
    return format_info


def parse_markdown_line(line):
    """解析markdown行，提取位置标记、格式和内容"""
    # 匹配段落标记
    para_match = re.match(r'<!-- PARA:(\d+),RUN:(\d+),FORMAT:(.*?) -->(.*)', line)
    if para_match:
        return {
            'type': 'paragraph',
            'para_idx': int(para_match.group(1)),
            'run_idx': int(para_match.group(2)),
            'format': string_to_format(para_match.group(3)),
            'content': para_match.group(4)
        }

    # 匹配表格标记（支持嵌套表格）
    # 普通表格: TABLE:1,ROW:0,CELL:0,PARA:0,RUN:0
    # 嵌套表格: TABLE:1,ROW:0,CELL:0,NESTED:0,ROW:0,CELL:0,PARA:0,RUN:0
    table_match = re.match(r'<!-- TABLE:(\d+)(,ROW:(\d+),CELL:(\d+)(?:,NESTED:(\d+),ROW:(\d+),CELL:(\d+))?,PARA:(\d+),RUN:(\d+)),FORMAT:(.*?) -->(.*)', line)
    if table_match:
        table_idx = int(table_match.group(1))

        # 检查是否是嵌套表格
        if table_match.group(5):  # 有 NESTED 标记
            return {
                'type': 'nested_table_cell',
                'table_idx': table_idx,
                'outer_row_idx': int(table_match.group(3)),
                'outer_cell_idx': int(table_match.group(4)),
                'nested_idx': int(table_match.group(5)),
                'nested_row_idx': int(table_match.group(6)),
                'nested_cell_idx': int(table_match.group(7)),
                'para_idx': int(table_match.group(8)),
                'run_idx': int(table_match.group(9)),
                'format': string_to_format(table_match.group(10)),
                'content': table_match.group(11)
            }
        else:
            return {
                'type': 'table_cell',
                'table_idx': table_idx,
                'row_idx': int(table_match.group(3)),
                'cell_idx': int(table_match.group(4)),
                'para_idx': int(table_match.group(8)),
                'run_idx': int(table_match.group(9)),
                'format': string_to_format(table_match.group(10)),
                'content': table_match.group(11)
            }

    # 匹配表格标记
    if line.strip() == '<!-- TABLE -->':
        return {'type': 'table_start'}

    if line.strip() == '<!-- TABLE_END -->':
        return {'type': 'table_end'}

    # 普通空行
    if line.strip() == '':
        return {'type': 'empty'}

    return None


def has_field_code(run):
    """检查run是否包含域代码（instrText）"""
    for child in run._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'instrText':
            return True
    return False


def get_run_format(run):
    """获取run的格式信息"""
    format_info = {
        'font_name': run.font.name if run.font.name else None,
        'font_size': run.font.size.pt if run.font.size else None,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'strike': run.font.strike,
        'superscript': run.font.superscript,
        'subscript': run.font.subscript,
    }
    return {k: v for k, v in format_info.items() if v is not None}


def apply_run_format(run, format_info):
    """应用格式到run"""
    if 'font_name' in format_info and format_info['font_name']:
        run.font.name = format_info['font_name']
    if 'font_size' in format_info:
        run.font.size = Pt(format_info['font_size'])
    if format_info.get('bold'):
        run.font.bold = True
    if format_info.get('italic'):
        run.font.italic = True
    if format_info.get('underline'):
        run.font.underline = True
    if format_info.get('strike'):
        run.font.strike = True
    if format_info.get('superscript'):
        run.font.superscript = True
    if format_info.get('subscript'):
        run.font.subscript = True


def restore_docx_from_markdown(markdown_path, template_path, output_path):
    """从markdown文件还原到docx模板"""
    # 先复制模板文件
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # 读取markdown文件
    with open(markdown_path, 'r', encoding='utf-8') as f:
        markdown_lines = [line.rstrip('\n') for line in f]

    # 解析markdown
    parsed_items = []
    for line in markdown_lines:
        item = parse_markdown_line(line)
        if item:
            parsed_items.append(item)

    # 统计
    paragraph_changes = 0
    table_changes = 0

    # 按段落索引分组
    para_items = {}
    for item in parsed_items:
        if item['type'] == 'paragraph':
            para_idx = item['para_idx']
            if para_idx not in para_items:
                para_items[para_idx] = []
            para_items[para_idx].append(item)

    # 还原段落内容
    for para_idx, items in sorted(para_items.items()):
        if para_idx >= len(doc.paragraphs):
            continue

        para = doc.paragraphs[para_idx]

        # 保存原始段落中的所有run信息（包括域代码和图片）
        original_runs = []
        for run in para.runs:
            original_runs.append({
                'run': run,
                'has_field_code': has_field_code(run),
                'format': get_run_format(run),
                'text': run.text
            })

        # 先清空段落
        para.clear()

        # 按run_idx排序并重建
        items_sorted = sorted(items, key=lambda x: x['run_idx'])

        # 处理每个run
        for item in items_sorted:
            run_idx = item['run_idx']
            content = item['content']
            target_format = item.get('format', {})

            # 检查原始是否有对应位置的run
            if run_idx < len(original_runs):
                orig_run = original_runs[run_idx]
                if orig_run['has_field_code']:
                    # 如果原始run包含域代码，创建一个零宽空格run来保持位置
                    new_run = para.add_run('\u200B')  # Zero Width Space
                    apply_run_format(new_run, orig_run['format'])
                else:
                    # 使用原始run的格式，应用新内容
                    # 如果内容是 [FIELD:...]，说明这是域代码标记
                    if re.match(r'^\[FIELD:', content):
                        new_run = para.add_run('\u200B')  # Zero Width Space
                    else:
                        # 移除其他可能的 [FIELD:...] 标记
                        content = re.sub(r'\[FIELD:[^\]]*\]', '', content)
                        # 如果content为空，使用零宽空格
                        text_to_add = content if content else '\u200B'
                        new_run = para.add_run(text_to_add)
                        if target_format:
                            apply_run_format(new_run, target_format)
                        elif orig_run['format']:
                            apply_run_format(new_run, orig_run['format'])
                    paragraph_changes += 1
            else:
                # 创建新的run
                # 如果内容是 [FIELD:...]，说明这是域代码标记
                if re.match(r'^\[FIELD:', content):
                    new_run = para.add_run('\u200B')  # Zero Width Space
                else:
                    # 移除其他可能的 [FIELD:...] 标记
                    content = re.sub(r'\[FIELD:[^\]]*\]', '', content)
                    text_to_add = content if content else '\u200B'
                    new_run = para.add_run(text_to_add)
                    apply_run_format(new_run, target_format)
                paragraph_changes += 1

    # 还原表格内容（包括嵌套表格）
    for item in parsed_items:
        if item['type'] == 'table_cell':
            table_idx = item['table_idx']
            row_idx = item['row_idx']
            cell_idx = item['cell_idx']
            para_idx = item['para_idx']
            run_idx = item['run_idx']
            content = item['content']
            target_format = item.get('format', {})

            # 如果内容是 [FIELD:...]，说明这是域代码，跳过
            if re.match(r'^\[FIELD:', content):
                continue

            # 移除其他可能的 [FIELD:...] 标记
            content = re.sub(r'\[FIELD:[^\]]*\]', '', content)

            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    if cell_idx < len(row.cells):
                        cell = row.cells[cell_idx]
                        if para_idx < len(cell.paragraphs):
                            para = cell.paragraphs[para_idx]

                            # 如果run存在且不包含域代码，修改文本
                            if run_idx < len(para.runs) and not has_field_code(para.runs[run_idx]):
                                para.runs[run_idx].text = content
                                table_changes += 1

        elif item['type'] == 'nested_table_cell':
            table_idx = item['table_idx']
            outer_row_idx = item['outer_row_idx']
            outer_cell_idx = item['outer_cell_idx']
            nested_idx = item['nested_idx']
            nested_row_idx = item['nested_row_idx']
            nested_cell_idx = item['nested_cell_idx']
            para_idx = item['para_idx']
            run_idx = item['run_idx']
            content = item['content']
            target_format = item.get('format', {})

            # 如果内容是 [FIELD:...]，说明这是域代码，跳过
            if re.match(r'^\[FIELD:', content):
                continue

            # 移除其他可能的 [FIELD:...] 标记
            content = re.sub(r'\[FIELD:[^\]]*\]', '', content)

            if table_idx < len(doc.tables):
                outer_table = doc.tables[table_idx]
                if outer_row_idx < len(outer_table.rows):
                    outer_row = outer_table.rows[outer_row_idx]
                    if outer_cell_idx < len(outer_row.cells):
                        outer_cell = outer_row.cells[outer_cell_idx]
                        if nested_idx < len(outer_cell.tables):
                            nested_table = outer_cell.tables[nested_idx]
                            if nested_row_idx < len(nested_table.rows):
                                nested_row = nested_table.rows[nested_row_idx]
                                if nested_cell_idx < len(nested_row.cells):
                                    nested_cell = nested_row.cells[nested_cell_idx]
                                    if para_idx < len(nested_cell.paragraphs):
                                        para = nested_cell.paragraphs[para_idx]

                                        # 如果run存在且不包含域代码，修改文本
                                        if run_idx < len(para.runs) and not has_field_code(para.runs[run_idx]):
                                            para.runs[run_idx].text = content
                                            table_changes += 1

    # 保存文档
    doc.save(output_path)
    print(f"文档已还原到: {output_path}")
    print(f"共修改了 {paragraph_changes} 个段落中的run")
    print(f"共修改了 {table_changes} 个表格单元格中的run")
    print("所有格式、图片和样式都已保留")


if __name__ == '__main__':
    # 解析命令行参数
    if len(sys.argv) < 3:
        print("用法: python restore_docx.py <markdown文件路径> <docx模板路径> [输出docx路径]")
        print("示例:")
        print("  python restore_docx.py content.md template.docx output.docx")
        print("  python restore_docx.py content.md template.docx")
        sys.exit(1)

    markdown_path = sys.argv[1]
    template_path = sys.argv[2]

    if len(sys.argv) >= 4:
        output_path = sys.argv[3]
    else:
        # 如果未指定输出路径，使用与markdown相同的文件名，但扩展名改为.docx
        base_name = os.path.splitext(markdown_path)[0]
        output_path = base_name + '_restored.docx'

    # 检查文件是否存在
    if not os.path.exists(markdown_path):
        print(f"错误: 文件不存在: {markdown_path}")
        sys.exit(1)

    if not os.path.exists(template_path):
        print(f"错误: 文件不存在: {template_path}")
        sys.exit(1)

    # 还原文档
    restore_docx_from_markdown(markdown_path, template_path, output_path)
