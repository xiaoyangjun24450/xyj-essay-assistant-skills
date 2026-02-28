#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取docx文件内容，按格式差异换行

用法:
    python extract_docx.py <docx文件路径> [输出文件路径]

示例:
    python extract_docx.py input.docx output.md
    python extract_docx.py input.docx  # 默认输出到stdout
"""
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_run_format_info(run):
    """获取run的格式信息"""
    if run is None:
        return None

    format_info = {
        'font_name': run.font.name if run.font.name else None,
        'font_size': run.font.size.pt if run.font.size else None,
        'bold': run.font.bold,
        'italic': run.font.italic,
        'underline': run.font.underline,
        'color': str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
        'strike': run.font.strike,
        'double_strike': run.font.double_strike,
        'superscript': run.font.superscript,
        'subscript': run.font.subscript,
        'all_caps': run.font.all_caps,
        'small_caps': run.font.small_caps,
    }

    return format_info


def formats_equal(format1, format2):
    """比较两个格式是否相同"""
    if format1 is None or format2 is None:
        return format1 == format2

    # 比较所有格式属性，但忽略font_name为None的情况
    for key in format1.keys():
        val1 = format1.get(key)
        val2 = format2.get(key)

        # 对于font_name，如果其中一个是None，则认为兼容
        if key == 'font_name':
            if val1 is not None and val2 is not None and val1 != val2:
                return False
        else:
            if val1 != val2:
                return False

    return True


def extract_runs_with_format(runs):
    """提取runs中的文本，按格式差异分割"""
    result = []
    current_format = None
    current_text = []

    for run in runs:
        if run.text == '':
            continue

        run_format = get_run_format_info(run)

        if current_format is None:
            current_format = run_format
            current_text.append(run.text)
        elif formats_equal(current_format, run_format):
            current_text.append(run.text)
        else:
            if current_text:
                result.append(''.join(current_text))
            current_format = run_format
            current_text = [run.text]

    if current_text:
        result.append(''.join(current_text))

    return result


def extract_table_cell_content(cell):
    """提取表格单元格的内容，每个格式块占一行"""
    parts = []
    for para in cell.paragraphs:
        run_parts = extract_runs_with_format(para.runs)
        for part in run_parts:
            if part:
                parts.append(part)
    return parts


def extract_docx_content(docx_path, output_path):
    """提取docx文件内容到txt文件，格式差异时换行"""
    doc = Document(docx_path)

    output_lines = []

    # 提取段落内容
    for para in doc.paragraphs:
        if not para.runs or all(run.text == '' for run in para.runs):
            output_lines.append('')
            continue

        parts = extract_runs_with_format(para.runs)
        for part in parts:
            output_lines.append(part)
        output_lines.append('')

    # 提取表格内容
    for table in doc.tables:
        # 表格标识
        output_lines.append('')
        output_lines.append('---')
        output_lines.append('')

        for row in table.rows:
            # 每个单元格的内容
            for cell in row.cells:
                cell_parts = extract_table_cell_content(cell)
                for part in cell_parts:
                    output_lines.append(part)
            # 行结束后换行
            output_lines.append('')

        output_lines.append('')
        output_lines.append('---')
        output_lines.append('')

    # 写入输出文件
    with open(output_path, 'w', encoding='utf-8') as f:
        for line in output_lines:
            f.write(line + '\n')

    print(f"内容已提取到: {output_path}")
    print(f"总行数: {len(output_lines)}")


if __name__ == '__main__':
    # 解析命令行参数
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <docx文件路径> [输出文件路径]")
        print("示例:")
        print("  python extract_docx.py input.docx output.md")
        print("  python extract_docx.py input.docx")
        sys.exit(1)

    docx_path = sys.argv[1]

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # 如果未指定输出路径，使用与输入相同的文件名，但扩展名改为.md
        base_name = os.path.splitext(docx_path)[0]
        output_path = base_name + '.md'

    # 检查输入文件是否存在
    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在: {docx_path}")
        sys.exit(1)

    # 提取内容
    extract_docx_content(docx_path, output_path)
